import docx
import io
import json
import os
import sys

# Ensure stdout uses UTF-8
sys.stdout.reconfigure(encoding='utf-8')

# Import docx2pdf safely
try:
    from docx2pdf import convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

# Mapped Teaching Plan Data (deterministic for Python course)
plan_data = {
    "curso": "Aperfeiçoamento Profissional - Programação em Python",
    "unidade_curricular": "Programação em Python",
    "carga_horaria_uc": "60h",
    "n_aulas": "15 (aulas de 4h)",
    "objetivo_uc": "Capacitar profissionais para desenvolver aplicações em linguagem Python, por meio de técnicas de programação, seguindo boas práticas, procedimentos e normas.",
    "situacoes_aprendizagem": [
        {
            "titulo": "SITUAÇÃO DE APRENDIZAGEM 01",
            "carga_horaria": "20h",
            "estrategia_tipo": "Situação-problema",
            "capacidades_tecnicas": [
                "Identificar os requisitos do problema para definição dos recursos a serem utilizados.",
                "Elaborar algoritmo da solução do problema.",
                "Configurar o ambiente de desenvolvimento em Python.",
                "Programar em linguagem Python."
            ],
            "capacidades_socioemocionais": [
                "Demonstrar autogestão.",
                "Demonstrar pensamento analítico."
            ],
            "conhecimentos": [
                "Fundamentos da computação (funcionamento de computadores, programas, algoritmos, linguagens de baixo/alto nível, interpretada/compilada).",
                "Lógica de programação (abstração lógica, álgebra booleana, fluxogramas, operadores aritméticos, relacionais, lógicos, expressões, teste de mesa, refatoração).",
                "Ambiente de programação em Python (histórico, instalação, configuração do ambiente, tipos de dados, variáveis, constantes, instruções de I/O, operadores, estruturas condicionais e de repetição)."
            ],
            "contextualizacao": "A empresa TechLogistics realiza manualmente a triagem e o cálculo de relatórios de envios diários de mercadorias. O processo consome cerca de 4 horas diárias de um assistente administrativo, o que atrasa a tomada de decisão logística e gera erros frequentes no fechamento financeiro do frete.",
            "observacoes_docente": "Recomenda-se dividir a turma em duplas para incentivar a colaboração analítica. Nas primeiras 4 horas, foque na modelagem do fluxograma antes da digitação do código. Oriente a instalação do VS Code e da versão mais recente do Python.",
            "desafio": "Como desenvolvedor júnior contratado pela TechLogistics, você deve configurar o ambiente Python e construir um script automatizado que leia um arquivo de texto com registros brutos de envios, faça o parse dos dados de frete, aplique taxas condicionais de distância e salve um sumário final limpo com a soma do faturamento total de frete.",
            "resultados_esperados": [
                "Fluxograma do algoritmo da solução (digital ou papel).",
                "Script Python (.py) estruturado para processamento automático do relatório.",
                "Arquivo de saída limpo contendo o sumário formatado."
            ],
            "anexos": "Arquivo bruto de exemplo (dados_envio.txt), fluxogramas de referência.",
            "referencias": "MENEZES, Nilo Ney Coutinho. Introdução à programação com Python. Novatec Editora, 2016.",
            "instrumento_registro": [
                {
                    "capacidade": "Identificar os requisitos do problema para definição dos recursos a serem utilizados",
                    "criterio": "[CRÍTICO] O aluno mapeou corretamente todas as variáveis e requisitos lógicos do problema de logística?"
                },
                {
                    "capacidade": "Elaborar algoritmo da solução do problema",
                    "criterio": "O fluxograma elaborado resolve a ordenação do processamento logicamente?"
                },
                {
                    "capacidade": "Configurar o ambiente de desenvolvimento em Python",
                    "criterio": "[CRÍTICO] O ambiente de programação (Python e VS Code) foi instalado e configurado corretamente?"
                },
                {
                    "capacidade": "Programar em linguagem Python",
                    "criterio": "[CRÍTICO] O script lê e processa as informações do arquivo de texto sem erros de sintaxe?"
                },
                {
                    "capacidade": "Demonstrar autogestão",
                    "criterio": "O aluno organizou seu tempo e entregou as etapas do projeto no prazo determinado?"
                },
                {
                    "capacidade": "Demonstrar pensamento analítico",
                    "criterio": "O aluno propôs soluções adequadas para contornar problemas de conversão de dados do relatório?"
                }
            ]
        },
        {
            "titulo": "SITUAÇÃO DE APRENDIZAGEM 02",
            "carga_horaria": "24h",
            "estrategia_tipo": "Projeto",
            "capacidades_tecnicas": [
                "Configurar o ambiente de desenvolvimento em Python (biblioteca Pygame).",
                "Programar em linguagem Python (funções, estruturas avançadas, manipulação de arquivos).",
                "Programar jogos 2D em linguagem Python."
            ],
            "capacidades_socioemocionais": [
                "Demonstrar inteligência emocional.",
                "Demonstrar autonomia."
            ],
            "conhecimentos": [
                "Programação avançada em Python (funções com/sem argumento/retorno, escopo de variáveis, exceções: try-except, manipulação de arquivos: leitura/gravação/exclusão).",
                "Fundamentos de programação gráfica 2D (Pygame, renderização de sprites, colisão, loop do jogo, interface do usuário [UI], animação)."
            ],
            "contextualizacao": "A startup educacional GameEdu deseja criar um mini-jogo interativo 2D em blocos (tipo Pyblock) para auxiliar crianças de 10 a 12 anos a entenderem conceitos espaciais. Eles precisam de um protótipo jogável com controles fluidos, salvamento de progresso e visual limpo.",
            "observacoes_docente": "Estimule os alunos a desenharem a lógica da interface do jogo (HUD) e o mapa de colisões antes de iniciarem a codificação gráfica. Ajude-os a entender o ciclo do Game Loop (Update, Draw, Tick).",
            "desafio": "Sua equipe foi contratada para desenvolver um protótipo de jogo 2D com a biblioteca Pygame. O jogo deve conter controle de personagem (teclado/mouse), detecção de colisão com objetos na tela, pontuação visível (UI), efeitos de animação e salvamento do recorde (High Score) em um arquivo local.",
            "resultados_esperados": [
                "Estrutura gráfica do jogo funcionando sem travamentos (código-fonte .py).",
                "Arquivo externo de recordes integrado ao fluxo do jogo.",
                "Manual de usuário básico de jogabilidade."
            ],
            "anexos": "Assets gráficos de exemplo (sprites, sons de fundo), rascunho de tela (mockup).",
            "referencias": "PyGame Documentation. Disponível em: https://www.pygame.org/news.",
            "instrumento_registro": [
                {
                    "capacidade": "Configurar o ambiente de desenvolvimento em Python",
                    "criterio": "[CRÍTICO] A biblioteca Pygame foi instalada e importada corretamente no script?"
                },
                {
                    "capacidade": "Programar em linguagem Python",
                    "criterio": "[CRÍTICO] O salvamento e leitura do recorde de pontuação em arquivo externo foi implementado com tratamento de exceções?"
                },
                {
                    "capacidade": "Programar jogos 2D em linguagem Python",
                    "criterio": "[CRÍTICO] O jogo implementa o ciclo do Game Loop, colisões e renderização gráfica de forma fluida?"
                },
                {
                    "capacidade": "Demonstrar inteligência emocional",
                    "criterio": "O aluno soube lidar com a frustração diante de bugs lógicos complexos durante a programação do loop?"
                },
                {
                    "capacidade": "Demonstrar autonomia",
                    "criterio": "O aluno buscou de forma autônoma documentações e referências extras para enriquecer a mecânica do jogo?"
                }
            ]
        },
        {
            "titulo": "SITUAÇÃO DE APRENDIZAGEM 03",
            "carga_horaria": "16h",
            "estrategia_tipo": "Situação-problema",
            "capacidades_tecnicas": [
                "Identificar os requisitos do problema para definição dos recursos a serem utilizados.",
                "Validar software por meio de testes."
            ],
            "capacidades_socioemocionais": [
                "Demonstrar pensamento analítico.",
                "Demonstrar autonomia."
            ],
            "conhecimentos": [
                "Processo para desenvolvimento de software (testes de software: testes unitários, testes de integração, testes de segurança).",
                "Validação de software (revisões com stakeholders, testes de aceitação de usuário [UAT], conformidade com requisitos, avaliação de usabilidade, verificação de necessidades do negócio)."
            ],
            "contextualizacao": "Durante os testes de usuário do jogo criado para a GameEdu, a equipe comercial identificou que scores negativos estavam sendo registrados e que arquivos de save corrompidos faziam o jogo travar na inicialização. A startup exige a validação formal do software por meio de testes estruturados antes de lançar o produto.",
            "observacoes_docente": "Apresente conceitos de Testes de Mesa versus Testes Unitários Automáticos. Ensine a usar o framework padrão `unittest` do Python. Foque na importância da cobertura de testes para caminhos alternativos de exceção.",
            "desafio": "Como Engenheiro de Qualidade de Software (QA), você deve construir uma suíte de testes unitários e de integração para o núcleo lógico do jogo anterior (cálculo de score e leitura de saves), assegurando que entradas inválidas sejam tratadas e gerando um relatório de validação.",
            "resultados_esperados": [
                "Scripts de teste unitário (test_game.py) usando o módulo unittest.",
                "Relatório de validação detalhando os cenários testados e os resultados (Pass/Fail)."
            ],
            "anexos": "Modelo de relatório de testes (tabelas de conformidade).",
            "referencias": "RELEASE, Python Tutorial. 3.10.4. https://docs.python.org/pt-br/3/tutorial/.",
            "instrumento_registro": [
                {
                    "capacidade": "Identificar os requisitos do problema para definição dos recursos a serem utilizados",
                    "criterio": "O aluno estruturou cenários de teste que cobrem todos os requisitos críticos do software?"
                },
                {
                    "capacidade": "Validar software por meio de testes",
                    "criterio": "[CRÍTICO] Os testes unitários criados validaram com sucesso as funções lógicas do jogo sob condições normais e anômalas?"
                },
                {
                    "capacidade": "Validar software por meio de testes",
                    "criterio": "[CRÍTICO] O aluno apresentou um relatório de validação detalhado mapeando conformidade com as regras de negócio?"
                },
                {
                    "capacidade": "Demonstrar pensamento analítico",
                    "criterio": "O aluno identificou brechas e falhas lógicas adicionais na lógica do software através da escrita de testes de borda?"
                }
            ]
        }
    ]
}

def fill_docx():
    template_path = r"c:\Users\Redes\Documents\GusPlan\MODELO PLANO DE ENSINO NOVO 2025 (1).docx"
    output_docx_path = r"c:\Users\Redes\Documents\GusPlan\Plano_de_Ensino_Python_MESEP.docx"
    output_pdf_path = r"c:\Users\Redes\Documents\GusPlan\Plano_de_Ensino_Python_MESEP.pdf"

    print("Carregando o modelo DOCX...")
    doc = docx.Document(template_path)

    curso = plan_data["curso"]
    uc = plan_data["unidade_curricular"]
    ch_uc = plan_data["carga_horaria_uc"]
    n_aulas = plan_data["n_aulas"]
    objetivo = plan_data["objetivo_uc"]
    sas = plan_data["situacoes_aprendizagem"]
    num_sas = len(sas)

    print("Preenchendo as Situações de Aprendizagem...")
    for i, sa in enumerate(sas):
        t_details_idx = i * 3
        t_ref_idx = i * 3 + 1
        t_inst_idx = i * 3 + 2

        if t_details_idx < len(doc.tables):
            t_details = doc.tables[t_details_idx]
            
            # Populating Details Table
            if len(t_details.rows) == 12:
                t_details.cell(1, 0).text = f"Curso: {curso}"
                t_details.cell(2, 0).text = f"Unidade curricular (UC): {uc}"
                t_details.cell(3, 0).text = f"Carga horária da UC: {ch_uc}"
                t_details.cell(3, 1).text = f"Nº de aulas: {n_aulas}"
                t_details.cell(4, 0).text = f"Carga horária prevista para o desenvolvimento da Situação de Aprendizagem: {sa['carga_horaria']}"
                t_details.cell(5, 0).text = f"Objetivo da UC: {objetivo}"

                # Checkbox de Capacidades
                has_tech = len(sa["capacidades_tecnicas"]) > 0
                has_socio = len(sa["capacidades_socioemocionais"]) > 0
                tech_mark = "X" if has_tech else " "
                socio_mark = "X" if has_socio else " "
                t_details.cell(6, 0).text = f"Capacidades a serem desenvolvidas:  Básicas (   )     Técnicas ( {tech_mark} )    Socioemocionais ( {socio_mark} )"

                # Capacidades
                cap_text = "Capacidades Técnicas:\n"
                for cap in sa["capacidades_tecnicas"]:
                    cap_text += f"- {cap}\n"
                cap_text += "\nCapacidades Socioemocionais:\n"
                for cap in sa["capacidades_socioemocionais"]:
                    cap_text += f"- {cap}\n"
                t_details.cell(7, 0).text = cap_text.strip()

                # Conhecimentos
                know_text = "Conhecimentos:\n"
                for kw in sa["conhecimentos"]:
                    know_text += f"- {kw}\n"
                t_details.cell(8, 0).text = know_text.strip()

                # Estratégia Desafiadora
                strat_type = sa["estrategia_tipo"]
                p_sp = "X" if strat_type == "Situação-problema" else " "
                p_ec = "X" if strat_type == "Estudo de caso" else " "
                p_pa = "X" if strat_type == "Pesquisa Aplicada" else " "
                p_pr = "X" if strat_type == "Projeto" else " "
                p_in = "X" if strat_type == "Integrador" else " "
                t_details.cell(9, 0).text = (
                    f"Estratégia de aprendizagem desafiadora\n"
                    f"Situação-problema ( {p_sp} )   Estudo de caso ( {p_ec} )  Pesquisa Aplicada ( {p_pa} )  "
                    f"Projeto ( {p_pr} ) Integrador ( {p_in} )"
                )

                # Contexto, Desafio, Resultados
                obs_docente = sa["observacoes_docente"]
                context_val = sa["contextualizacao"]
                desafio_val = sa["desafio"]
                res_esperados = "\n".join([f"- {r}" for r in sa["resultados_esperados"]])
                
                t_details.cell(10, 0).text = (
                    f"Contextualização:\n{context_val}\n\n"
                    f"Observações para o docente:\n{obs_docente}\n\n"
                    f"Desafio:\n{desafio_val}\n\n"
                    f"Resultados esperados:\n{res_esperados}"
                )

                t_details.cell(11, 0).text = f"Anexos (Figuras, esquemas, desenhos, leiaute, formulários, etc):\n{sa['anexos']}"

            elif len(t_details.rows) == 3:
                # Subsequent SA details table
                strat_type = sa["estrategia_tipo"]
                p_sp = "X" if strat_type == "Situação-problema" else " "
                p_ec = "X" if strat_type == "Estudo de caso" else " "
                p_pa = "X" if strat_type == "Pesquisa Aplicada" else " "
                p_pr = "X" if strat_type == "Projeto" else " "
                p_in = "X" if strat_type == "Integrador" else " "
                t_details.cell(0, 0).text = (
                    f"Estratégia de aprendizagem desafiadora\n"
                    f"Situação-problema ( {p_sp} )   Estudo de caso ( {p_ec} )  Pesquisa Aplicada ( {p_pa} )  "
                    f"Projeto ( {p_pr} ) Integrador ( {p_in} )"
                )

                # Contexto, Desafio, Resultados
                obs_docente = sa["observacoes_docente"]
                context_val = sa["contextualizacao"]
                desafio_val = sa["desafio"]
                res_esperados = "\n".join([f"- {r}" for r in sa["resultados_esperados"]])
                t_details.cell(1, 0).text = (
                    f"Contextualização:\n{context_val}\n\n"
                    f"Observações para o docente:\n{obs_docente}\n\n"
                    f"Desafio:\n{desafio_val}\n\n"
                    f"Resultados esperados:\n{res_esperados}"
                )

                t_details.cell(2, 0).text = f"Anexos (Figuras, esquemas, desenhos, leiaute, formulários, etc):\n{sa['anexos']}"

        if t_ref_idx < len(doc.tables):
            t_ref = doc.tables[t_ref_idx]
            t_ref.cell(0, 0).text = f"Referências (livros, apostilas, sites, blog etc.):\n{sa['referencias']}"

        if t_inst_idx < len(doc.tables):
            t_inst = doc.tables[t_inst_idx]
            instrument_data = sa["instrumento_registro"]
            
            for idx, pair in enumerate(instrument_data):
                row_idx = 4 + idx
                if row_idx < len(t_inst.rows):
                    t_inst.cell(row_idx, 0).text = pair["capacidade"]
                    
                    crit_cell = t_inst.cell(row_idx, 1)
                    crit_cell.text = ""
                    p = crit_cell.paragraphs[0]
                    criterio_text = pair["criterio"]
                    
                    if "[CRÍTICO]" in criterio_text:
                        clean_crit = criterio_text.replace("[CRÍTICO]", "").strip()
                        run = p.add_run(clean_crit)
                        run.bold = True
                    else:
                        p.add_run(criterio_text)

            # Deletar linhas vazias na tabela de registro
            rows_to_keep = 4 + len(instrument_data)
            if rows_to_keep < len(t_inst.rows):
                rows_to_delete = sorted(list(range(rows_to_keep, len(t_inst.rows))), reverse=True)
                for r_idx in rows_to_delete:
                    if r_idx >= 4:
                        tr = t_inst.rows[r_idx]._tr
                        t_inst._tbl.remove(tr)

    print("Removendo Situações de Aprendizagem não utilizadas (SA 4)...")
    total_sas_in_template = 4
    if num_sas < total_sas_in_template:
        tables_to_delete = []
        for sa_idx in range(num_sas, total_sas_in_template):
            tables_to_delete.extend([sa_idx * 3, sa_idx * 3 + 1, sa_idx * 3 + 2])
        
        tables_to_delete.sort(reverse=True)
        for t_idx in tables_to_delete:
            if t_idx < len(doc.tables):
                tbl = doc.tables[t_idx]
                tbl._tbl.getparent().remove(tbl._tbl)

        vale_lembrar_count = 0
        paras_to_remove = []
        for para in list(doc.paragraphs):
            text = para.text.strip()
            if "Vale lembrar que:" in text:
                vale_lembrar_count += 1
                if vale_lembrar_count > num_sas:
                    paras_to_remove.append(para)
                    continue
            if vale_lembrar_count > num_sas:
                boilerplates = [
                    "as estratégias e recursos",
                    "a intervenção mediadora",
                    "a avaliação formativa é feita",
                    "Legenda:",
                    "A= Atingiu",
                    "N= Não atingiu",
                    "Negrito = Crítico",
                    "Sem negrito = Desejável"
                ]
                if any(b in text for b in boilerplates) or text == "":
                    paras_to_remove.append(para)
        
        for para in paras_to_remove:
            try:
                para._p.getparent().remove(para._p)
            except Exception:
                pass

    print("Preenchendo Tabela de Níveis de Desempenho...")
    if len(doc.tables) > 0:
        t_grades = doc.tables[-1]
        if "TABELA DE NÍVEIS DE DESEMPENHO" in t_grades.cell(0, 0).text:
            t_grades.cell(2, 0).text = "Todos os critérios Críticos e ≥ 90% dos critérios Desejáveis atingidos."
            t_grades.cell(2, 1).text = "Atingiu com Excelência (A)"
            t_grades.cell(2, 2).text = "9,0 - 10,0"
            
            t_grades.cell(3, 0).text = "Todos os critérios Críticos e 70% a 89% dos critérios Desejáveis atingidos."
            t_grades.cell(3, 1).text = "Atingiu Plenamente (B)"
            t_grades.cell(3, 2).text = "8,0 - 8,9"
            
            t_grades.cell(4, 0).text = "Todos os critérios Críticos e 50% a 69% dos critérios Desejáveis atingidos."
            t_grades.cell(4, 1).text = "Atingiu Parcialmente (C)"
            t_grades.cell(4, 2).text = "7,0 - 7,9"
            
            t_grades.cell(5, 0).text = "Um ou mais critérios Críticos não atingidos e/ou < 50% dos critérios Desejáveis atingidos."
            t_grades.cell(5, 1).text = "Não Atingiu (D) - Retenção/Recuperação"
            t_grades.cell(5, 2).text = "0,0 - 6,9"
            
            if len(t_grades.rows) > 6:
                rows_to_delete = sorted(list(range(6, len(t_grades.rows))), reverse=True)
                for r_idx in rows_to_delete:
                    tr = t_grades.rows[r_idx]._tr
                    t_grades._tbl.remove(tr)

    print(f"Salvando arquivo DOCX em: {output_docx_path}")
    doc.save(output_docx_path)
    print("Salvo com sucesso!")

    if HAS_DOCX2PDF:
        try:
            print("Convertendo arquivo DOCX para PDF (requer Microsoft Word)...")
            convert(output_docx_path, output_pdf_path)
            print(f"PDF salvo com sucesso em: {output_pdf_path}")
        except Exception as e:
            print(f"Não foi possível converter para PDF automaticamente: {e}")
            print("Você pode abrir o arquivo DOCX gerado e salvá-lo como PDF no próprio Word.")
    else:
        print("Biblioteca docx2pdf não encontrada. Apenas o DOCX foi gerado.")

if __name__ == "__main__":
    fill_docx()
