from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import pypdf
import docx
import io
import json
import os
import urllib.request
import urllib.error

app = FastAPI()

# Enable CORS for local testing and Vercel hosting
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
@app.get("/api")
def read_root():
    return {"status": "GusPlan API is running successfully"}

@app.post("/api/extract")
async def extract_text(file: UploadFile = File(...)):
    filename = file.filename.lower()
    content = await file.read()
    
    # Guard: reject files over 15MB to avoid timeout
    if len(content) > 15 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Arquivo muito grande (máximo 15MB). Para o livro MESEP, use a versão comprimida.")
    
    if filename.endswith(".pdf"):
        try:
            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file)
            text = ""
            total_pages = len(reader.pages)
            # Limit to first 180 pages to avoid serverless timeout on large books
            max_pages = min(total_pages, 180)
            for i, page in enumerate(reader.pages):
                if i >= max_pages:
                    break
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if total_pages > max_pages:
                text += f"\n[Nota: Documento com {total_pages} páginas. Apenas as primeiras {max_pages} foram carregadas.]"
            return {"text": text.strip(), "pages_total": total_pages, "pages_loaded": max_pages}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Erro ao extrair PDF: {str(e)}")
            
    elif filename.endswith(".docx"):
        try:
            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)
            text = ""
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            return {"text": text.strip()}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Erro ao extrair DOCX: {str(e)}")
            
    else:
        raise HTTPException(status_code=400, detail="Formato de arquivo não suportado. Envie PDF ou DOCX.")

# Model fallback order — tries each until one works
GEMINI_MODELS_FALLBACK = [
    "gemini-2.5-flash-lite-preview-06-17",
    "gemini-2.5-flash",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
]

def call_gemini_api(api_key: str, model: str, contents: list, system_instruction: str = None, response_mime_type: str = None):
    """Call Gemini REST API with automatic fallback to other models on quota errors."""
    # Build the candidate model list: preferred model first, then fallbacks
    candidate_models = [model] + [m for m in GEMINI_MODELS_FALLBACK if m != model]
    
    last_error = None
    for try_model in candidate_models:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{try_model}:generateContent?key={api_key}"
        
        payload = {"contents": contents}
        
        if system_instruction:
            payload["systemInstruction"] = {"parts": [{"text": system_instruction}]}
            
        generation_config = {}
        if response_mime_type:
            generation_config["responseMimeType"] = response_mime_type
        if generation_config:
            payload["generationConfig"] = generation_config
            
        req = urllib.request.Request(
            url,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        
        try:
            with urllib.request.urlopen(req) as response:
                res_data = json.loads(response.read().decode("utf-8"))
                return res_data['candidates'][0]['content']['parts'][0]['text']
        except urllib.error.HTTPError as e:
            err_body = e.read().decode("utf-8")
            # If quota error (429) or model not found (404), try next model
            if e.code in (429, 404):
                last_error = f"Modelo {try_model} indisponível (HTTP {e.code}). Tentando próximo..."
                continue
            raise Exception(f"Erro HTTP {e.code} da API Gemini ({try_model}): {err_body}")
        except Exception as e:
            last_error = str(e)
            continue
    
    raise Exception(f"Todos os modelos Gemini falharam. Último erro: {last_error}")

def get_api_key(passed_key: str = None):
    key = passed_key or os.environ.get("GEMINI_API_KEY", "")
    if not key:
        return ""
    # Try base64 decoding in case Vercel env validators blocked the raw key
    try:
        import base64
        padded_key = key
        if len(key) % 4 != 0:
            padded_key += "=" * (4 - len(key) % 4)
        decoded = base64.b64decode(padded_key).decode("utf-8")
        if decoded.startswith("AQ.") or decoded.startswith("AIza"):
            return decoded
    except Exception:
        pass
    return key

@app.post("/api/chat")
async def chat_with_docs(
    message: str = Form(...),
    history: str = Form("[]"),  # JSON string
    course_plan: str = Form(""),
    mesep_book: str = Form(""),
    api_key: str = Form(None),
    model_name: str = Form(None)
):
    try:
        active_key = get_api_key(api_key)
        if not active_key:
            raise HTTPException(status_code=400, detail="API Key do Gemini não configurada no servidor.")
            
        active_model = model_name or os.environ.get("GEMINI_MODEL", "gemini-2.0-flash-lite")
        history_list = json.loads(history)
        
        # Build system prompt/context
        system_instruction = (
            "Você é o GusPlan, um assistente especialista em planos de ensino do SENAI baseado na Metodologia MESEP.\n"
            "Você tem acesso aos seguintes documentos enviados pelo usuário:\n\n"
            f"--- PLANO DE CURSO ---\n{course_plan[:10000]}\n\n"
            f"--- LIVRO DIRETRIZES MESEP (TRECHO) ---\n{mesep_book[:15000]}\n\n"
            "Responda às dúvidas do professor de forma concisa, profissional e focada nos padrões de qualidade do SENAI. "
            "Seja amigável e ofereça conselhos de como estruturar Situações de Aprendizagem (SAs) desafiadoras, critérios de avaliação "
            "e instrumentos de registro. Escreva em Português do Brasil."
        )
        
        # Format chat history for Gemini (standard REST structure)
        contents = []
        for msg in history_list:
            role = "user" if msg.get("role") == "user" else "model"
            contents.append({"role": role, "parts": [{"text": msg.get("text", "")}]})
            
        # Append current message
        contents.append({"role": "user", "parts": [{"text": f"Contexto do Sistema:\n{system_instruction}\n\nPergunta do usuário:\n{message}"}]})
        
        response_text = call_gemini_api(active_key, active_model, contents)
        return {"response": response_text}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro no chat com Gemini: {str(e)}")

@app.post("/api/analyze")
async def analyze_course(
    course_plan: str = Form(...),
    mesep_book: str = Form(""),
    api_key: str = Form(None),
    model_name: str = Form(None)
):
    try:
        active_key = get_api_key(api_key)
        if not active_key:
            raise HTTPException(status_code=400, detail="API Key do Gemini não configurada no servidor.")
            
        active_model = model_name or os.environ.get("GEMINI_MODEL", "gemini-2.0-flash-lite")
        genai.configure(api_key=active_key)
        model = genai.GenerativeModel(active_model)
        
        prompt = (
            "Você é o GusPlan, um assistente de Inteligência Artificial especializado na Metodologia SENAI de Educação Profissional (MESEP).\n"
            "Sua tarefa é analisar o PLANO DE CURSO fornecido abaixo e estruturar um plano de ensino completo adaptado ao MESEP SENAI.\n\n"
            f"--- PLANO DE CURSO ---\n{course_plan}\n\n"
            "Diretrizes importantes do MESEP SENAI:\n"
            "1. Divida a carga horária total da Unidade Curricular em Situações de Aprendizagem (SAs) lógicas e coerentes. "
            "Para cursos pequenos (até 60h), 2 a 3 SAs são ideais. Para cursos maiores, até 4 SAs.\n"
            "2. Cada SA deve ter uma estratégia de aprendizagem desafiadora (como Situação-problema ou Projeto).\n"
            "3. O 'Desafio' deve descrever o papel do aluno (ex: 'Você é um desenvolvedor... e deve...') e o problema de maneira estimulante.\n"
            "4. A 'Contextualização' deve criar um cenário real e envolvente de mercado de trabalho.\n"
            "5. Os 'Critérios de Avaliação' na tabela de instrumentos devem ser formulados como perguntas objetivas. "
            "Eles devem ser classificados como Críticos (essenciais para a ocupação) ou Desejáveis. "
            "IMPORTANTE: Adicione o prefixo [CRÍTICO] no início da pergunta dos critérios críticos. Ex: '[CRÍTICO] O aluno configurou o ambiente de desenvolvimento corretamente?'. "
            "Os critérios sem o prefixo [CRÍTICO] serão considerados desejáveis.\n"
            "6. Distribua as Capacidades Técnicas e Socioemocionais e os Conhecimentos listados no plano de curso de forma lógica entre as SAs criadas.\n\n"
            "Retorne a resposta estritamente no formato JSON, sem marcações markdown de código (como ```json) ou qualquer outro texto explicativo fora do JSON. "
            "O JSON deve seguir exatamente a seguinte estrutura:\n\n"
            "{\n"
            '  "curso": "Nome do Curso (ex: Aperfeiçoamento Profissional - Programação em Python)",\n'
            '  "unidade_curricular": "Nome da Unidade Curricular",\n'
            '  "carga_horaria_uc": "Carga Horária Total (ex: 60h)",\n'
            '  "n_aulas": "Número aproximado de aulas (ex: 15)",\n'
            '  "objetivo_uc": "Objetivo geral da UC extraído do plano de curso",\n'
            '  "situacoes_aprendizagem": [\n'
            "    {\n"
            '      "titulo": "SITUAÇÃO DE APRENDIZAGEM 01",\n'
            '      "carga_horaria": "Carga horária prevista para esta SA (ex: 20h)",\n'
            '      "estrategia_tipo": "Situação-problema",\n' # Escolha entre: Situação-problema, Estudo de caso, Pesquisa Aplicada, Projeto, Integrador
            '      "capacidades_tecnicas": ["Capacidade Técnica 1", "Capacidade Técnica 2"],\n'
            '      "capacidades_socioemocionais": ["Capacidade Socioemocional 1", "Capacidade Socioemocional 2"],\n'
            '      "conhecimentos": ["Conhecimento 1", "Conhecimento 2"],\n'
            '      "contextualizacao": "Texto de contextualização conectando o aluno ao mercado...",\n'
            '      "observacoes_docente": "Dicas e orientações pedagógicas para o docente...",\n'
            '      "desafio": "Texto detalhado do desafio prático...",\n'
            '      "resultados_esperados": ["Resultado esperado 1", "Resultado esperado 2"],\n'
            '      "anexos": "Anexos e leiaute recomendados...",\n'
            '      "referencias": "Referências recomendadas de livros ou sites...",\n'
            '      "instrumento_registro": [\n'
            "        {\n"
            '          "capacidade": "Capacidade que está sendo avaliada",\n'
            '          "criterio": "[CRÍTICO] Pergunta do critério de avaliação ou Pergunta de critério desejável"\n'
            "        }\n"
            "      ]\n"
            "    }\n"
            "  ]\n"
            "}"
        )
        
        response = model.generate_content(
            prompt,
            generation_config={"response_mime_type": "application/json"}
        )
        
        # Parse JSON to confirm it is valid
        parsed_data = json.loads(response.text)
        return parsed_data
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro na análise com Gemini: {str(e)}")

@app.post("/api/generate")
async def generate_docx(
    template: UploadFile = File(...),
    data: str = Form(...)  # JSON string of teaching plan
):
    try:
        plan_data = json.loads(data)
        template_content = await template.read()
        
        # Open the template document
        doc = docx.Document(io.BytesIO(template_content))
        
        curso = plan_data.get("curso", "")
        uc = plan_data.get("unidade_curricular", "")
        ch_uc = plan_data.get("carga_horaria_uc", "")
        n_aulas = plan_data.get("n_aulas", "")
        objetivo = plan_data.get("objetivo_uc", "")
        sas = plan_data.get("situacoes_aprendizagem", [])
        num_sas = len(sas)
        
        # Populate each Situation of Learning
        for i, sa in enumerate(sas):
            if i >= 4: # Template supports up to 4 SAs out of the box
                break
                
            # Table index map
            t_details_idx = i * 3
            t_ref_idx = i * 3 + 1
            t_inst_idx = i * 3 + 2
            
            if t_details_idx < len(doc.tables):
                t_details = doc.tables[t_details_idx]
                
                # Populating Details Table
                if len(t_details.rows) == 12:
                    t_details.cell(1, 0).text = f"Curso: {curso}"
                    t_details.cell(2, 0).text = f"Unidade curricular (UC): {uc}"
                    t_details.cell(3, 0).text = f"Carga horária da UC: {ch_uc}"
                    t_details.cell(3, 1).text = f"Nº de aulas: {n_aulas}"
                    t_details.cell(4, 0).text = f"Carga horária prevista para o desenvolvimento da Situação de Aprendizagem: {sa.get('carga_horaria', '')}"
                    t_details.cell(5, 0).text = f"Objetivo da UC: {objetivo}"
                    
                    # Checkbox selection for Strategy Type
                    has_tech = len(sa.get("capacidades_tecnicas", [])) > 0
                    has_socio = len(sa.get("capacidades_socioemocionais", [])) > 0
                    tech_mark = "X" if has_tech else " "
                    socio_mark = "X" if has_socio else " "
                    t_details.cell(6, 0).text = f"Capacidades a serem desenvolvidas:  Básicas (   )     Técnicas ( {tech_mark} )    Socioemocionais ( {socio_mark} )"
                    
                    # Row 7: list capacities
                    cap_text = "Capacidades Técnicas:\n"
                    for cap in sa.get("capacidades_tecnicas", []):
                        cap_text += f"- {cap}\n"
                    cap_text += "\nCapacidades Socioemocionais:\n"
                    for cap in sa.get("capacidades_socioemocionais", []):
                        cap_text += f"- {cap}\n"
                    t_details.cell(7, 0).text = cap_text.strip()
                    
                    # Row 8: Conhecimentos
                    know_text = "Conhecimentos:\n"
                    for kw in sa.get("conhecimentos", []):
                        know_text += f"- {kw}\n"
                    t_details.cell(8, 0).text = know_text.strip()
                    
                    # Row 9: Challenging Strategy Type Checkbox
                    strat_type = sa.get("estrategia_tipo", "")
                    p_sp = "X" if strat_type == "Situação-problema" else " "
                    p_ec = "X" if strat_type == "Estudo de caso" else " "
                    p_pa = "X" if strat_type == "Pesquisa Aplicada" else " "
                    p_pr = "X" if strat_type == "Projeto" else " "
                    p_in = "X" if strat_type == "Integrador" else " "
                    t_details.cell(9, 0).text = (
                        f"Estratégia de aprendizagem desafiadora\n"
                        f"Situação-problema ( {p_sp} )   Estudo de caso ( {p_ec} )  Pesquisa Aplicada ( {p_pa} )  "
                        f"Projeto ( {p_pr} ) Integrador ( {p_in} )"
                    )
                    
                    # Row 10: Context, Teacher Notes, Desafio, Resultados
                    obs_docente = sa.get("observacoes_docente", "")
                    context_val = sa.get("contextualizacao", "")
                    desafio_val = sa.get("desafio", "")
                    res_esperados = "\n".join([f"- {r}" for r in sa.get("resultados_esperados", [])])
                    t_details.cell(10, 0).text = (
                        f"Contextualização:\n{context_val}\n\n"
                        f"Observações para o docente:\n{obs_docente}\n\n"
                        f"Desafio:\n{desafio_val}\n\n"
                        f"Resultados esperados:\n{res_esperados}"
                    )
                    
                    # Row 11: Anexos
                    t_details.cell(11, 0).text = f"Anexos (Figuras, esquemas, desenhos, leiaute, formulários, etc):\n{sa.get('anexos', '')}"
                
                elif len(t_details.rows) == 3:
                    # Subsequent SA details table
                    # Row 0: Strategy Selection Checkbox
                    strat_type = sa.get("estrategia_tipo", "")
                    p_sp = "X" if strat_type == "Situação-problema" else " "
                    p_ec = "X" if strat_type == "Estudo de caso" else " "
                    p_pa = "X" if strat_type == "Pesquisa Aplicada" else " "
                    p_pr = "X" if strat_type == "Projeto" else " "
                    p_in = "X" if strat_type == "Integrador" else " "
                    t_details.cell(0, 0).text = (
                        f"Estratégia de aprendizagem desafiadora\n"
                        f"Situação-problema ( {p_sp} )   Estudo de caso ( {p_ec} )  Pesquisa Aplicada ( {p_pa} )  "
                        f"Projeto ( {p_pr} ) Integrador ( {p_in} )"
                    )
                    
                    # Row 1: Context, Teacher Notes, Desafio, Resultados
                    obs_docente = sa.get("observacoes_docente", "")
                    context_val = sa.get("contextualizacao", "")
                    desafio_val = sa.get("desafio", "")
                    res_esperados = "\n".join([f"- {r}" for r in sa.get("resultados_esperados", [])])
                    t_details.cell(1, 0).text = (
                        f"Contextualização:\n{context_val}\n\n"
                        f"Observações para o docente:\n{obs_docente}\n\n"
                        f"Desafio:\n{desafio_val}\n\n"
                        f"Resultados esperados:\n{res_esperados}"
                    )
                    
                    # Row 2: Anexos
                    t_details.cell(2, 0).text = f"Anexos (Figuras, esquemas, desenhos, leiaute, formulários, etc):\n{sa.get('anexos', '')}"
                
            # Populate References Table
            if t_ref_idx < len(doc.tables):
                t_ref = doc.tables[t_ref_idx]
                t_ref.cell(0, 0).text = f"Referências (livros, apostilas, sites, blog etc.):\n{sa.get('referencias', '')}"
                
            # Populate Instrument Table (Table 2)
            if t_inst_idx < len(doc.tables):
                t_inst = doc.tables[t_inst_idx]
                instrument_data = sa.get("instrumento_registro", [])
                
                # Row 4 to 19 can be populated.
                # Table 2 has 20 rows total.
                for idx, pair in enumerate(instrument_data):
                    row_idx = 4 + idx
                    if row_idx < len(t_inst.rows):
                        t_inst.cell(row_idx, 0).text = pair.get("capacidade", "")
                        
                        # Populate Criterion and format [CRÍTICO] in bold
                        crit_cell = t_inst.cell(row_idx, 1)
                        crit_cell.text = "" # Clear text
                        p = crit_cell.paragraphs[0]
                        criterio_text = pair.get("criterio", "")
                        
                        if "[CRÍTICO]" in criterio_text:
                            clean_crit = criterio_text.replace("[CRÍTICO]", "").strip()
                            run = p.add_run(clean_crit)
                            run.bold = True
                        else:
                            p.add_run(criterio_text)
                            
                # Delete empty unused rows in the instrument table
                rows_to_keep = 4 + len(instrument_data)
                # Keep at least a few rows or exactly the number of criteria
                if rows_to_keep < len(t_inst.rows):
                    rows_to_delete = sorted(list(range(rows_to_keep, len(t_inst.rows))), reverse=True)
                    for r_idx in rows_to_delete:
                        # Don't delete headers (rows 0-3)
                        if r_idx >= 4:
                            tr = t_inst.rows[r_idx]._tr
                            t_inst._tbl.remove(tr)

        # Remove Unused SAs from document if we generated less than 4
        # Total SAs in template is 4 (Tables 0-11)
        total_sas_in_template = 4
        if num_sas < total_sas_in_template:
            # Delete tables from index num_sas * 3 up to 11 in reverse order
            tables_to_delete = []
            for sa_idx in range(num_sas, total_sas_in_template):
                tables_to_delete.extend([sa_idx * 3, sa_idx * 3 + 1, sa_idx * 3 + 2])
            
            tables_to_delete.sort(reverse=True)
            for t_idx in tables_to_delete:
                if t_idx < len(doc.tables):
                    tbl = doc.tables[t_idx]
                    tbl._tbl.getparent().remove(tbl._tbl)
            
            # Clean up trailing boilerplate paragraphs for unused SAs
            vale_lembrar_count = 0
            paras_to_remove = []
            for para in list(doc.paragraphs):
                text = para.text.strip()
                if "Vale lembrar que:" in text:
                    vale_lembrar_count += 1
                    if vale_lembrar_count > num_sas:
                        paras_to_remove.append(para)
                        continue
                if vale_lembrar_count > num_sas:
                    boilerplates = [
                        "as estratégias e recursos",
                        "a intervenção mediadora",
                        "a avaliação formativa é feita",
                        "Legenda:",
                        "A= Atingiu",
                        "N= Não atingiu",
                        "Negrito = Crítico",
                        "Sem negrito = Desejável"
                    ]
                    if any(b in text for b in boilerplates) or text == "":
                        paras_to_remove.append(para)
            
            for para in paras_to_remove:
                try:
                    para._p.getparent().remove(para._p)
                except Exception:
                    pass

        # Populate Table 12 (Tabela de Níveis de Desempenho) which is now the last table
        if len(doc.tables) > 0:
            t_grades = doc.tables[-1]
            # Ensure it is the grading table
            if "TABELA DE NÍVEIS DE DESEMPENHO" in t_grades.cell(0, 0).text:
                # Let's populate rows 2, 3, 4, 5 with standard values
                t_grades.cell(2, 0).text = "Todos os critérios Críticos e ≥ 90% dos critérios Desejáveis atingidos."
                t_grades.cell(2, 1).text = "Atingiu com Excelência (A)"
                t_grades.cell(2, 2).text = "9,0 - 10,0"
                
                t_grades.cell(3, 0).text = "Todos os critérios Críticos e 70% a 89% dos critérios Desejáveis atingidos."
                t_grades.cell(3, 1).text = "Atingiu Plenamente (B)"
                t_grades.cell(3, 2).text = "8,0 - 8,9"
                
                t_grades.cell(4, 0).text = "Todos os critérios Críticos e 50% a 69% dos critérios Desejáveis atingidos."
                t_grades.cell(4, 1).text = "Atingiu Parcialmente (C)"
                t_grades.cell(4, 2).text = "7,0 - 7,9"
                
                t_grades.cell(5, 0).text = "Um ou mais critérios Críticos não atingidos e/ou < 50% dos critérios Desejáveis atingidos."
                t_grades.cell(5, 1).text = "Não Atingiu (D) - Retenção/Recuperação"
                t_grades.cell(5, 2).text = "0,0 - 6,9"
                
                # Delete unused rows in Table 12 (rows 6, 7, 8)
                if len(t_grades.rows) > 6:
                    rows_to_delete = sorted(list(range(6, len(t_grades.rows))), reverse=True)
                    for r_idx in rows_to_delete:
                        tr = t_grades.rows[r_idx]._tr
                        t_grades._tbl.remove(tr)

        # Save to buffer
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        
        return StreamingResponse(
            output_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Plano_de_Ensino_MESEP.docx"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao gerar arquivo DOCX: {str(e)}")
